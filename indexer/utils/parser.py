import os
import csv
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Optional
from urllib.parse import urlparse

import PyPDF2
from docx import Document
from PIL import Image
import requests
from bs4 import BeautifulSoup
from google import genai
from google.genai import types
from dotenv import load_dotenv

load_dotenv()


class BaseParser(ABC):

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)

    @abstractmethod
    def parse(self, source: str) -> List[Dict[str, str]]:
        pass

    def _create_document(self, content: str, metadata: Optional[Dict] = None) -> Dict[str, str]:
        doc = {"content": content}
        if metadata:
            doc["metadata"] = metadata
        return doc


class TextParser(BaseParser):

    def parse(self, source: str) -> List[Dict[str, str]]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"Text file not found: {source}")

        try:
            with open(source, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = {
                "source": source,
                "type": "text",
                "filename": os.path.basename(source)
            }

            return [self._create_document(content, metadata)]
        except Exception as e:
            self.logger.error(f"Error parsing text file {source}: {e}")
            raise


class PDFParser(BaseParser):

    def __init__(self):
        super().__init__()

    def parse(self, source: str) -> List[Dict[str, str]]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"PDF file not found: {source}")

        documents = []

        try:
            with open(source, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    content = page.extract_text()

                    if content.strip():
                        metadata = {
                            "source": source,
                            "type": "pdf",
                            "filename": os.path.basename(source),
                            "page": page_num,
                            "total_pages": total_pages
                        }
                        documents.append(self._create_document(content, metadata))

            self.logger.info(f"Parsed {len(documents)} pages from PDF: {source}")
            return documents
        except Exception as e:
            self.logger.error(f"Error parsing PDF {source}: {e}")
            raise


class DocxParser(BaseParser):

    def __init__(self):
        super().__init__()

    def parse(self, source: str) -> List[Dict[str, str]]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"DOCX file not found: {source}")

        try:
            doc = Document(source)

            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)

            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)

            if table_texts:
                content += "\n\n" + "\n".join(table_texts)

            metadata = {
                "source": source,
                "type": "docx",
                "filename": os.path.basename(source),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables)
            }

            return [self._create_document(content, metadata)]
        except Exception as e:
            self.logger.error(f"Error parsing DOCX file {source}: {e}")
            raise


class ImageOCRParser(BaseParser):

    def __init__(self):
        super().__init__()
        api_key = os.getenv("GEMINI_API")
        if not api_key:
            raise ValueError("GEMINI_API not found in environment variables")
        self.client = genai.Client(api_key=api_key)

    def parse(self, source: str) -> List[Dict[str, str]]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"Image file not found: {source}")

        try:
            image = Image.open(source)
            image_format = image.format or "PNG"

            with open(source, "rb") as f:
                image_bytes = f.read()

            mime_type = f"image/{image_format.lower()}"
            if image_format.upper() == "JPG":
                mime_type = "image/jpeg"

            prompt = """Analyze this image and provide a detailed transcription of all text and UI elements.

This image is likely a UI screenshot. Please:
1. Extract ALL visible text exactly as shown
2. Identify and describe key UI elements (buttons, input fields, menus, icons, labels, headers)
3. Note the layout and positioning of elements (e.g., "top navigation bar", "sidebar menu", "main content area")
4. Describe any visual indicators (colors, highlights, selected states, error messages)
5. Capture any data shown in tables, lists, or cards

Format your response as a structured description that preserves the semantic meaning and organization of the UI."""

            response = self.client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    types.Part.from_text(text=prompt),
                    types.Part.from_bytes(data=image_bytes, mime_type=mime_type)
                ]
            )
            content = response.text

            metadata = {
                "source": source,
                "type": "image_ocr",
                "filename": os.path.basename(source),
                "image_format": image_format,
                "image_size": f"{image.width}x{image.height}"
            }

            return [self._create_document(content, metadata)]
        except Exception as e:
            self.logger.error(f"Error parsing image {source}: {e}")
            raise


class CSVParser(BaseParser):

    def parse(self, source: str, delimiter: str = ',', has_header: bool = True) -> List[Dict[str, str]]:
        if not os.path.exists(source):
            raise FileNotFoundError(f"CSV file not found: {source}")

        documents = []

        try:
            with open(source, 'r', encoding='utf-8') as file:
                csv_reader = csv.DictReader(file, delimiter=delimiter) if has_header else csv.reader(file, delimiter=delimiter)

                for row_num, row in enumerate(csv_reader, start=1):
                    if has_header:
                        content = "\n".join(f"{key}: {value}" for key, value in row.items() if value)
                    else:
                        content = ", ".join(str(cell) for cell in row)

                    metadata = {
                        "source": source,
                        "type": "csv",
                        "filename": os.path.basename(source),
                        "row": row_num
                    }
                    documents.append(self._create_document(content, metadata))

            self.logger.info(f"Parsed {len(documents)} rows from CSV: {source}")
            return documents
        except Exception as e:
            self.logger.error(f"Error parsing CSV {source}: {e}")
            raise


class URLParser(BaseParser):

    def __init__(self, timeout: int = 10):
        super().__init__()
        self.timeout = timeout
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })

    def parse(self, source: str) -> List[Dict[str, str]]:
        try:
            parsed = urlparse(source)
            if not parsed.scheme or not parsed.netloc:
                raise ValueError(f"Invalid URL: {source}")

            response = self.session.get(source, timeout=self.timeout)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            for script in soup(["script", "style"]):
                script.decompose()

            content = soup.get_text(separator='\n', strip=True)

            title = soup.find('title')
            title_text = title.get_text() if title else None

            metadata = {
                "source": source,
                "type": "url",
                "url": source,
                "title": title_text,
                "status_code": response.status_code
            }

            return [self._create_document(content, metadata)]
        except Exception as e:
            self.logger.error(f"Error parsing URL {source}: {e}")
            raise
